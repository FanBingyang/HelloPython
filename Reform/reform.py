# !/usr/bin/env python
# -*- encoding: utf-8 -*-
'''
@File   :  reform.py    
@Desc   :  对中文进行翻译处理，输出和原文意思一样但语序不一样的句子
@Author :  ByFan
@Time   :  2021/6/3 10:10 
'''

import requests
import random
import json
from hashlib import md5
import docx
from tkinter import filedialog, Tk

APP_ID = "20210602000851704"
# APP_ID = "1234"
APP_KEY = "29Dsw3no3Xl70zyefoa6"

endpoint = 'http://api.fanyi.baidu.com'
path = '/api/trans/vip/translate'
url = endpoint + path

# 定于语种对应值
CHINESE = 'zh'      # 中文
ENGLISH = 'en'      # 英语
FRENCH = 'fra'      # 法语
GERMAN = 'de'       # 德语


headers = {'Content-Type': 'application/x-www-form-urlencoded'}


"""
    生成签名
"""
def make_md5(s,encoding='utf-8'):
    return md5(s.encode(encoding)).hexdigest()

"""
    自动识别语言并转换成指定语种，也可指定原始语言
"""
def translate(query,to_lang,from_lang='auto'):
    salt = random.randint(32768, 65536)     # 随机数
    sign = make_md5(APP_ID + query + str(salt) + APP_KEY)   # 构建md5签名
    # 封装请求参数
    payload = {'appid': APP_ID,
               'q': query,
               'from': from_lang,
               'to': to_lang,
               'salt': salt,
               'sign': sign
               }
    # 发起请求
    r = requests.post(url, params=payload, headers=headers)
    request_json = r.json()
    # print(json.dumps(request_json, indent=4, ensure_ascii=False))
    result = ''
    if 'error_code' in request_json:
        return '错误，请联系管理员排查'

    # 返回翻译过后的文字
    return request_json['trans_result'][0]['dst']

"""
    经过几个语种之间的翻译，进行改进
"""
def reform(text):
    english = translate(text, ENGLISH)
    german = translate(english, GERMAN)
    french = translate(german, FRENCH)
    chinese = translate(french, CHINESE)
    return chinese

"""
    选择文件
"""
def choose_file():
    root = Tk()
    root.withdraw()
    file_path = (filedialog.askopenfilename(title='选择文档',filetypes=[('All Files','docx')]))
    return file_path

"""
    删除文档中的元素（段落、表格等等）
"""
def delete_element(DocxElement):
    p = DocxElement._element
    p.getparent().remove(p)
    # p._p = p._element = None  # 与下面语句等同
    DocxElement._p = DocxElement._element = None


"""
    对文档内容进行修改
"""
def reDocx(fileName):
    # 打开文件
    file = docx.Document(fileName)
    print("文档段落总数为：" + str(len(file.paragraphs)))

    file.add_paragraph("\n\n以下是修改过后的文章：\n")

    # 循环对文档的每个段落进行修改
    for para in file.paragraphs:
        text = para.text.strip()
        if text != "":
            text = reform(text)
        file.add_paragraph(text)
        # print("处理中。。。"+str(i+1)+"/"+str(len(file.paragraphs)))

    # 因为在对文档操作之前添加了一句修改说明，所以在操作的时候也把这句话给操作了
    # 操作完成之后要把当前文档最后一个段落（也就是添加的那一句修改说明）给删除了
    delete_element(file.paragraphs[len(file.paragraphs)-1])

    file.save(fileName)

if __name__ == "__main__":
    while True:
        print("选择操作类型：文字输入还是文档\n1、文字\t2、文档\t0、退出")
        switch = input("输入对应数字：")
        if(switch == '1'):
            while True:
               inText = input("输入文字（输入0退出）：")
               # 去除两端空格
               text = inText.strip()
               if text == '0':
                   print("退出文字操作!\n")
                   break
               elif text == "":
                   print("输入不能为空")
               else :
                   print("正在处理。。。。。。")
                   result = reform(text)
                   print("修改之后：",result,"\n")
        elif(switch == '2'):
            # fileName = input("输入docx文件名称：")
            # fileName = fileName + '.docx'
            print("选择文件：")
            # 选择文件
            fileName = choose_file()
            print(fileName)
            print("正在处理。。。。。。")
            # 对文档进行操作
            reDocx(fileName)
            print("操作结束，内容保存在原文档中！\n")
        elif(switch == '0'):
            print("关闭操作！\n")
            break
        else:
            print("请输入正确的操作数字！\n")


