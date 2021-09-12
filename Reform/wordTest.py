# !/usr/bin/env python
# -*- encoding: utf-8 -*-
'''
@File   :  wordTest.py    
@Desc   :  
@Author :  ByFan
@Time   :  2021/6/3 23:56 
'''
import docx
# from docx.shared import Cm
# 删除文档中的元素
def delete_paragraph(DocxElement):
    p = DocxElement._element
    p.getparent().remove(p)
    # p._p = p._element = None  # 与下面语句等同
    DocxElement._p = DocxElement._element = None

if __name__ == "__main__":
    print("Hello Word!")
    file = docx.Document("测试文档.docx")
    print("段落数："+str(len(file.paragraphs)))

    for para in file.paragraphs:
        text = para.text.strip()
        if text == "":
            text = "这是空行的代替"
        file.add_paragraph(text)

        # print(para.text)

    # delete_paragraph(file.tables[0])

    # file.add_paragraph("再添加一个图片")
    # 添加图片
    # file.add_picture('2.jpg',width=docx.shared.Cm(16),height=docx.shared.Cm(12))
    file.save("测试文档.docx")
